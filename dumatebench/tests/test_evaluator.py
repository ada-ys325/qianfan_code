import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from dumatebench.evaluator import (
    evaluate_contain,
    evaluate_directory_structure,
    evaluate_docx_font_style,
    evaluate_docx_paragraph_format,
    evaluate_excel_cell_number_tolerance,
    evaluate_excel_cell_style,
    evaluate_excel_formula,
    evaluate_excel_sheet_exists,
    evaluate_file_format_valid,
    evaluate_files_unchanged,
    evaluate_log_budget,
    evaluate_no_extra_files,
    evaluate_not_contain,
    evaluate_no_unexpected_diff,
    run_checklist_score,
)


try:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill
except ImportError:
    openpyxl = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    Document = None


class EvaluatorTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)

    def tearDown(self):
        self.tmp.cleanup()

    def test_file_format_valid_for_common_lightweight_types(self):
        (self.root / "answer.txt").write_text("hello")
        (self.root / "data.json").write_text(json.dumps({"ok": True}))
        (self.root / "brief.pdf").write_bytes(b"%PDF-1.4\n%minimal\n")

        self.assertTrue(evaluate_file_format_valid(self.root, {"file": "answer.txt"}))
        self.assertTrue(evaluate_file_format_valid(self.root, {"file": "data.json"}))
        self.assertTrue(evaluate_file_format_valid(self.root, {"file": "brief.pdf"}))

        (self.root / "bad.json").write_text("{bad")
        self.assertFalse(evaluate_file_format_valid(self.root, {"file": "bad.json"}))

    def test_content_checks_support_markdown_json_html_svg_and_zip(self):
        (self.root / "answer.md").write_text("# Project Alpha\nDone")
        (self.root / "data.json").write_text(json.dumps({"project": "Project Alpha", "done": True}))
        (self.root / "page.html").write_text("<h1>Project Alpha</h1><script>ignored()</script>")
        (self.root / "drawing.svg").write_text('<svg><text id="title">Project Alpha</text></svg>')
        with zipfile.ZipFile(self.root / "submission.zip", "w") as archive:
            archive.writestr("src/main.py", "# Project Alpha")

        for file_name, doc_type in (
            ("answer.md", "md"),
            ("data.json", "json"),
            ("page.html", "html"),
            ("drawing.svg", "svg"),
            ("submission.zip", "zip"),
        ):
            self.assertTrue(
                evaluate_contain(
                    self.root,
                    {"file": file_name, "doc_type": doc_type, "keywords": ["Project Alpha"]},
                )
            )
            self.assertTrue(
                evaluate_file_format_valid(self.root, {"file": file_name, "doc_type": doc_type})
            )

    def test_not_contain_does_not_pass_for_missing_file(self):
        self.assertFalse(
            evaluate_not_contain(
                self.root,
                {"file": "missing.md", "doc_type": "md", "keywords": ["forbidden"]},
            )
        )

    def test_checklist_marks_unsupported_and_runtime_errors_ineligible(self):
        (self.root / "answer.md").write_text("complete")
        result = run_checklist_score(
            self.root,
            [
                {
                    "id": "valid",
                    "type": "evaluate_contain",
                    "args": {"file": "answer.md", "doc_type": "md", "keywords": ["complete"]},
                },
                {
                    "id": "unsupported",
                    "type": "evaluate_contain",
                    "args": {"file": "answer.md", "doc_type": "unknown", "keywords": ["complete"]},
                },
                {"id": "unknown_function", "type": "does_not_exist", "args": {}},
            ],
        )
        self.assertEqual(
            [item["status"] for item in result["checks"]],
            ["pass", "unsupported", "unsupported"],
        )
        self.assertEqual(result["partial_pass"], 1.0)
        self.assertEqual(result["eligible_check_count"], 1)
        self.assertEqual(result["ineligible_check_count"], 2)

    def test_files_unchanged_checks_reference_bytes(self):
        (self.root / "result.txt").write_text("do not change")
        (self.root / "reference.txt").write_text("do not change")

        args = {"matches": [{"file": "result.txt", "reference_file": "reference.txt"}]}
        self.assertTrue(evaluate_files_unchanged(self.root, args))

        (self.root / "result.txt").write_text("changed")
        self.assertFalse(evaluate_files_unchanged(self.root, args))

    def test_no_unexpected_diff_allows_only_matching_changes(self):
        (self.root / "before.txt").write_text("Name: Alice\nStatus: Draft\n")
        (self.root / "after.txt").write_text("Name: Alice\nStatus: Final\n")
        args = {
            "input_file": "before.txt",
            "output_file": "after.txt",
            "allowed_patterns": [r"^Status: "],
        }
        self.assertTrue(evaluate_no_unexpected_diff(self.root, args))

        (self.root / "after.txt").write_text("Name: Bob\nStatus: Final\n")
        self.assertFalse(evaluate_no_unexpected_diff(self.root, args))

    def test_directory_structure_and_no_extra_files(self):
        (self.root / "deliverables").mkdir()
        (self.root / "deliverables" / "report.txt").write_text("report")
        (self.root / "deliverables" / "notes.tmp").write_text("ignore")

        self.assertTrue(
            evaluate_directory_structure(
                self.root,
                {
                    # Checklist paths are already relative to the evaluator
                    # testbed; ``root`` is retained only for compatibility.
                    "root": "/",
                    "required_files": ["deliverables/report.txt"],
                    "required_dirs": [],
                    "forbidden_paths": ["deliverables/old.txt"],
                },
            )
        )
        self.assertTrue(
            evaluate_no_extra_files(
                self.root,
                {
                    "root": "deliverables",
                    "expected_files": ["report.txt"],
                    "ignore_patterns": ["*.tmp"],
                },
            )
        )

        (self.root / "deliverables" / "unexpected.txt").write_text("extra")
        self.assertFalse(
            evaluate_no_extra_files(
                self.root,
                {"root": "deliverables", "expected_files": ["report.txt"], "ignore_patterns": []},
            )
        )

    @unittest.skipIf(openpyxl is None, "openpyxl is not installed")
    def test_excel_sheet_formula_and_numeric_tolerance(self):
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Summary"
        sheet["A1"] = 100
        sheet["A2"] = 105
        sheet["A3"] = "=SUM(A1:A2)"
        workbook.create_sheet("RawData")
        workbook.save(self.root / "metrics.xlsx")
        workbook.close()

        self.assertTrue(
            evaluate_excel_sheet_exists(
                self.root,
                {"file": "metrics.xlsx", "required_sheets": ["Summary", "RawData"]},
            )
        )
        self.assertTrue(
            evaluate_excel_formula(
                self.root,
                {"file": "metrics.xlsx", "matches": [{"row": 3, "col": 1, "formula": "=SUM(A1:A2)"}]},
            )
        )
        self.assertTrue(
            evaluate_excel_cell_number_tolerance(
                self.root,
                {
                    "file": "metrics.xlsx",
                    "matches": [{"row": 1, "col": 1, "value": 100.1, "tolerance": 0.2}],
                },
            )
        )
        self.assertFalse(
            evaluate_excel_cell_number_tolerance(
                self.root,
                {
                    "file": "metrics.xlsx",
                    "matches": [{"row": 1, "col": 1, "value": 101, "tolerance": 0.2}],
                },
            )
        )

    @unittest.skipIf(openpyxl is None, "openpyxl is not installed")
    def test_excel_cell_style(self):
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Summary"
        cell = sheet["B2"]
        cell.value = 42
        cell.font = Font(name="Arial", size=14, bold=True, color="FF112233")
        cell.fill = PatternFill(fill_type="solid", fgColor="FFCCEEFF")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.number_format = "0.00"
        workbook.save(self.root / "styled.xlsx")
        workbook.close()

        self.assertTrue(
            evaluate_excel_cell_style(
                self.root,
                {
                    "file": "styled.xlsx",
                    "matches": [
                        {
                            "sheet": "Summary",
                            "row": 2,
                            "col": 2,
                            "font_name": "Arial",
                            "font_size_pt": 14,
                            "bold": True,
                            "font_color_rgb": "112233",
                            "fill_color_rgb": "CCEEFF",
                            "number_format": "0.00",
                            "horizontal_alignment": "center",
                            "vertical_alignment": "center",
                        }
                    ],
                },
            )
        )
        self.assertFalse(
            evaluate_excel_cell_style(
                self.root,
                {
                    "file": "styled.xlsx",
                    "matches": [{"sheet": "Summary", "row": 2, "col": 2, "bold": False}],
                },
            )
        )

    @unittest.skipIf(Document is None, "python-docx is not installed")
    def test_docx_font_style_and_paragraph_format(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.style = document.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run("Styled heading")
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.italic = False
        run.font.color.rgb = RGBColor(0x11, 0x22, 0x33)
        document.save(self.root / "styled.docx")

        self.assertTrue(
            evaluate_docx_font_style(
                self.root,
                {
                    "file": "styled.docx",
                    "matches": [
                        {
                            "paragraph_number": 1,
                            "run_text": "Styled",
                            "font_name": "Arial",
                            "font_size_pt": 16,
                            "bold": True,
                            "italic": False,
                            "color_rgb": "112233",
                        }
                    ],
                },
            )
        )
        self.assertTrue(
            evaluate_docx_paragraph_format(
                self.root,
                {
                    "file": "styled.docx",
                    "matches": [
                        {
                            "text": "Styled heading",
                            "alignment": "CENTER",
                            "left_indent_pt": 18,
                            "space_before_pt": 6,
                            "space_after_pt": 12,
                            "line_spacing": 1.5,
                            "style_name": "Normal",
                        }
                    ],
                },
            )
        )
        self.assertFalse(
            evaluate_docx_paragraph_format(
                self.root,
                {"file": "styled.docx", "matches": [{"text": "Styled heading", "alignment": "LEFT"}]},
            )
        )

    def test_log_budget_accepts_json_and_text_logs(self):
        (self.root / "run.json").write_text(json.dumps({"total_tokens": 1200, "elapsed_seconds": 33.5}))
        self.assertTrue(
            evaluate_log_budget(
                self.root,
                {"log_file": "run.json", "max_tokens": 1500, "max_time_seconds": 40},
            )
        )
        self.assertFalse(
            evaluate_log_budget(
                self.root,
                {"log_file": "run.json", "max_tokens": 1000, "max_time_seconds": 40},
            )
        )

        (self.root / "run.log").write_text("total_tokens: 88\nelapsed_seconds: 9.5\n")
        self.assertTrue(
            evaluate_log_budget(
                self.root,
                {"log_file": "run.log", "max_tokens": 100, "max_time_seconds": 10},
            )
        )


if __name__ == "__main__":
    unittest.main()
