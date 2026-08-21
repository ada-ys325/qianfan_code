from __future__ import annotations

import json
import mimetypes
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import quote

TEXT_SUFFIXES = {".txt", ".md", ".json", ".docx", ".html", ".htm", ".yaml", ".yml"}
MEDIA_MIME_TYPES = {
    ".mp3": "audio/mpeg", ".wav": "audio/wav", ".m4a": "audio/mp4",
    ".flac": "audio/flac", ".aac": "audio/aac", ".ogg": "audio/ogg",
    ".mp4": "video/mp4", ".mov": "video/quicktime", ".webm": "video/webm",
    ".mkv": "video/x-matroska",
}
IGNORED_FILENAMES = {"reward.json", "judge_result.json", "rubric.json"}


class ArtifactError(ValueError):
    pass


@dataclass(frozen=True)
class MediaConfig:
    mode: str = "data_url"
    max_bytes: int = 20 * 1024 * 1024
    base_url: str | None = None
    video_mode: str = "frames"
    video_frame_count: int = 8
    video_frame_max_bytes: int = 2 * 1024 * 1024
    ffmpeg_path: str = "ffmpeg"
    ffprobe_path: str = "ffprobe"

    @classmethod
    def from_env(cls) -> "MediaConfig":
        integer_env = {
            "max_bytes": ("DU_MATE_MEDIA_MAX_BYTES", 20 * 1024 * 1024),
            "video_frame_count": ("DU_MATE_VIDEO_FRAME_COUNT", 8),
            "video_frame_max_bytes": ("DU_MATE_VIDEO_FRAME_MAX_BYTES", 2 * 1024 * 1024),
        }
        values: dict[str, int] = {}
        for field, (name, default) in integer_env.items():
            try:
                values[field] = int(os.environ.get(name, str(default)))
            except ValueError as exc:
                raise ArtifactError(f"{name} must be an integer") from exc
        config = cls(
            mode=os.environ.get("DU_MATE_MEDIA_MODE", "data_url"),
            base_url=os.environ.get("DU_MATE_MEDIA_BASE_URL") or None,
            video_mode=os.environ.get("DU_MATE_VIDEO_MODE", "frames"),
            ffmpeg_path=os.environ.get("DU_MATE_FFMPEG_PATH", "ffmpeg"),
            ffprobe_path=os.environ.get("DU_MATE_FFPROBE_PATH", "ffprobe"),
            **values,
        )
        config.validate()
        return config

    def validate(self) -> None:
        if self.mode not in {"data_url", "url", "disabled"}:
            raise ArtifactError("media mode must be data_url, url, or disabled")
        if self.max_bytes <= 0:
            raise ArtifactError("media max_bytes must be positive")
        if self.video_mode not in {"frames", "video_url"}:
            raise ArtifactError("video mode must be frames or video_url")
        if not 1 <= self.video_frame_count <= 32:
            raise ArtifactError("video frame count must be between 1 and 32")
        if self.video_frame_max_bytes <= 0:
            raise ArtifactError("video frame max bytes must be positive")
        if self.mode == "url" and not self.base_url:
            raise ArtifactError("DU_MATE_MEDIA_BASE_URL is required when media mode is url")


def _sample_text(text: str, limit: int) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    head = int(limit * 0.55)
    middle = int(limit * 0.2)
    tail = limit - head - middle
    middle_start = max(head, len(text) // 2 - middle // 2)
    sampled = "\n...[sampled]...\n".join((text[:head], text[middle_start:middle_start + middle], text[-tail:]))
    return sampled, True


def _media_transport(path: Path, relative_path: str, config: MediaConfig) -> dict[str, Any]:
    size = path.stat().st_size
    if config.mode == "disabled":
        return {"status": "cannot_assess", "mode": "disabled", "reason": "media transport is disabled"}
    if size > config.max_bytes:
        return {
            "status": "cannot_assess", "mode": config.mode,
            "reason": f"media size {size} exceeds configured limit {config.max_bytes}",
        }
    if config.mode == "url":
        url = f"{config.base_url.rstrip('/')}/{quote(relative_path, safe='/')}"
        return {"status": "ready", "mode": "url", "reference": url}
    return {"status": "ready", "mode": "data_url", "reference": str(path)}


def _docx_text(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ArtifactError("python-docx is required to inspect .docx files") from exc
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
    text = "\n".join(paragraphs + table_cells)
    return text, {"paragraph_count": len(paragraphs), "table_count": len(document.tables),
                  "character_count": len(text)}


def _is_ignored(path: Path) -> bool:
    name = path.name.lower()
    return name in IGNORED_FILENAMES or (
        path.suffix.lower() == ".json" and name.startswith(("rubric", "judge_", "reward"))
    )


def extract_artifact(
    path: Path, *, root: Path, char_limit: int = 80_000,
    media_config: MediaConfig | None = None,
) -> dict[str, Any]:
    root = root.resolve()
    path = path.resolve()
    if root != path and root not in path.parents:
        raise ArtifactError(f"artifact escapes root: {path}")
    if not path.is_file():
        raise ArtifactError(f"artifact is not a file: {path}")
    relative_path = path.relative_to(root).as_posix()
    suffix = path.suffix.lower()
    size = path.stat().st_size
    if suffix in MEDIA_MIME_TYPES:
        config = media_config or MediaConfig.from_env()
        config.validate()
        mime_type = MEDIA_MIME_TYPES[suffix]
        return {
            "path": relative_path, "source_path": str(path), "suffix": suffix,
            "mime_type": mime_type, "size_bytes": size,
            "category": mime_type.split("/", 1)[0], "content": None,
            "truncated": False, "structure": {},
            "transport": _media_transport(path, relative_path, config),
        }
    if suffix not in TEXT_SUFFIXES:
        raise ArtifactError(f"unsupported artifact type: {suffix or '<none>'}")
    if suffix == ".docx":
        text, structure = _docx_text(path)
    else:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            raise ArtifactError(f"text artifact is not valid UTF-8: {relative_path}") from exc
        structure = {"character_count": len(text)}
        if suffix == ".json":
            try:
                value = json.loads(text)
            except json.JSONDecodeError as exc:
                raise ArtifactError(f"invalid JSON artifact: {relative_path}") from exc
            structure.update({"root_type": "object" if isinstance(value, dict) else "array" if isinstance(value, list) else type(value).__name__})
            if isinstance(value, dict):
                structure["top_level_keys"] = sorted(map(str, value.keys()))[:100]
            elif isinstance(value, list):
                structure["item_count"] = len(value)
    sampled, truncated = _sample_text(text, char_limit)
    mime_type = mimetypes.guess_type(path.name)[0] or "text/plain"
    return {
        "path": relative_path, "source_path": str(path), "suffix": suffix,
        "mime_type": mime_type, "size_bytes": size, "category": "text",
        "content": sampled, "truncated": truncated,
        "structure": structure, "transport": None,
    }


def collect_artifacts(
    root: Path, *, max_files: int = 20, total_chars: int = 60_000,
    media_config: MediaConfig | None = None,
) -> list[dict[str, Any]]:
    root = root.resolve()
    if not root.exists():
        return []
    supported = TEXT_SUFFIXES | set(MEDIA_MIME_TYPES)
    candidates = sorted(
        path for path in root.rglob("*")
        if path.is_file() and not _is_ignored(path)
        and not any(part.startswith(".") for part in path.relative_to(root).parts)
        and path.suffix.lower() in supported
    )
    artifacts: list[dict[str, Any]] = []
    remaining_chars = total_chars
    for path in candidates[:max_files]:
        resolved = path.resolve()
        if root != resolved and root not in resolved.parents:
            continue
        limit = max(1, remaining_chars) if path.suffix.lower() in TEXT_SUFFIXES else 1
        artifact = extract_artifact(path, root=root, char_limit=limit, media_config=media_config)
        artifacts.append(artifact)
        if artifact["category"] == "text":
            remaining_chars -= len(artifact["content"] or "")
    return artifacts


def artifact_inventory(artifacts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    inventory = []
    for item in artifacts:
        entry = {
            "path": item["path"], "mime_type": item["mime_type"],
            "size_bytes": item["size_bytes"], "category": item["category"],
            "truncated": item["truncated"], "structure": item["structure"],
        }
        if item["category"] == "text":
            entry["excerpt"] = item["content"]
        else:
            entry["transport_status"] = item["transport"]["status"]
            if item["transport"]["status"] != "ready":
                entry["cannot_assess_reason"] = item["transport"]["reason"]
        inventory.append(entry)
    return inventory
