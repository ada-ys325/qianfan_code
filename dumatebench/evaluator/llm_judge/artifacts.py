from __future__ import annotations

import json
from collections import Counter
from pathlib import Path
from typing import Any


SUPPORTED_SUFFIXES = {".txt", ".md", ".json", ".docx", ".html", ".htm"}
IGNORED_FILENAMES = {"reward.json", "judge_result.json", "rubric.json"}


class ArtifactError(ValueError):
    pass


def _sample_text(text: str, limit: int) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    head = int(limit * 0.55)
    middle = int(limit * 0.2)
    tail = limit - head - middle
    middle_start = max(head, len(text) // 2 - middle // 2)
    sampled = (
        text[:head]
        + f"\n\n[...中间省略 {middle_start - head} 字符...]\n\n"
        + text[middle_start : middle_start + middle]
        + f"\n\n[...尾部前省略 {len(text) - (middle_start + middle) - tail} 字符...]\n\n"
        + text[-tail:]
    )
    return sampled, True


def _plain_text(path: Path) -> tuple[str, dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return text, {"line_count": len(text.splitlines()), "character_count": len(text)}


def _json_text(path: Path) -> tuple[str, dict[str, Any]]:
    raw = path.read_text(encoding="utf-8", errors="strict")
    data = json.loads(raw)
    text = json.dumps(data, ensure_ascii=False, indent=2, sort_keys=True)
    if isinstance(data, dict):
        shape = {"root_type": "object", "top_level_keys": list(data)[:100]}
    elif isinstance(data, list):
        shape = {"root_type": "array", "item_count": len(data)}
    else:
        shape = {"root_type": type(data).__name__}
    shape["character_count"] = len(text)
    return text, shape


def _docx_text(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ArtifactError("python-docx is required to inspect .docx files") from exc

    document = Document(path)
    paragraphs: list[str] = []
    style_counts: Counter[str] = Counter()
    alignments: Counter[str] = Counter()
    font_names: Counter[str] = Counter()
    font_sizes: Counter[str] = Counter()
    non_black_runs = 0
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text
        if text:
            paragraphs.append(f"[P{index}] {text}")
        style_counts[str(paragraph.style.name if paragraph.style else "<none>")] += 1
        alignments[str(paragraph.alignment or "INHERITED")] += 1
        for run in paragraph.runs:
            if run.font.name:
                font_names[run.font.name] += 1
            if run.font.size:
                font_sizes[f"{run.font.size.pt:.1f}pt"] += 1
            color = run.font.color.rgb
            if color is not None and str(color).upper() != "000000":
                non_black_runs += 1

    tables: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.replace("\n", " ").strip() for cell in row.cells))
        tables.append(f"[TABLE {table_index}]\n" + "\n".join(rows))

    header_footer: list[str] = []
    for section_index, section in enumerate(document.sections, start=1):
        header = "\n".join(p.text for p in section.header.paragraphs if p.text)
        footer = "\n".join(p.text for p in section.footer.paragraphs if p.text)
        if header:
            header_footer.append(f"[HEADER {section_index}] {header}")
        if footer:
            header_footer.append(f"[FOOTER {section_index}] {footer}")

    text = "\n".join(header_footer + paragraphs + tables)
    summary = {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "style_counts": dict(style_counts.most_common(20)),
        "alignment_counts": dict(alignments.most_common()),
        "font_names": dict(font_names.most_common(20)),
        "font_sizes": dict(font_sizes.most_common(20)),
        "non_black_explicit_font_runs": non_black_runs,
        "character_count": len(text),
    }
    return text, summary


def extract_artifact(path: Path, *, root: Path, char_limit: int = 80000) -> dict[str, Any]:
    root = root.resolve()
    path = path.resolve()
    if root != path and root not in path.parents:
        raise ArtifactError(f"artifact escapes root: {path}")
    if not path.is_file():
        raise ArtifactError(f"artifact is not a file: {path}")
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ArtifactError(f"unsupported textual artifact type: {suffix or '<none>'}")

    if suffix == ".json":
        text, structure = _json_text(path)
        kind = "json"
    elif suffix == ".docx":
        text, structure = _docx_text(path)
        kind = "docx"
    else:
        text, structure = _plain_text(path)
        kind = suffix.lstrip(".")
    sampled, truncated = _sample_text(text, char_limit)
    return {
        "path": path.relative_to(root).as_posix(),
        "kind": kind,
        "size_bytes": path.stat().st_size,
        "truncated": truncated,
        "structure": structure,
        "content": sampled,
    }


def collect_artifacts(
    root: Path,
    *,
    max_files: int = 20,
    char_limit_per_file: int = 80000,
    total_char_limit: int = 240000,
) -> list[dict[str, Any]]:
    root = root.resolve()
    if not root.is_dir():
        raise ArtifactError(f"artifact root is not a directory: {root}")
    paths = sorted(
        path for path in root.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_SUFFIXES
        and path.name not in IGNORED_FILENAMES
        and not any(part.startswith(".") or part == "uploads_raw" for part in path.relative_to(root).parts)
    )
    if len(paths) > max_files:
        paths = paths[:max_files]

    artifacts: list[dict[str, Any]] = []
    remaining = total_char_limit
    for path in paths:
        if remaining <= 0:
            break
        artifact = extract_artifact(path, root=root, char_limit=min(char_limit_per_file, remaining))
        remaining -= len(artifact["content"])
        artifacts.append(artifact)
    return artifacts


def artifact_inventory(artifacts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "path": item["path"],
            "kind": item["kind"],
            "size_bytes": item["size_bytes"],
            "truncated": item["truncated"],
            "structure": item["structure"],
        }
        for item in artifacts
    ]
